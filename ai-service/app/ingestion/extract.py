"""Extractors for PDF/DOCX/TXT/MD (ADR-011). One function per format, always
dispatched by the format Java already validated by magic bytes — never by
trusting a client-supplied extension alone (.claude/rules/security.md).
"""

import re
from pathlib import Path

import pdfplumber
from docx import Document as DocxDocument

from app.models.ingestion import ExtractedBlock, ExtractedDocument

# Matches "1. Foo", "1.2 Bar", "1.2.3 Baz" — the numbered-section convention
# used throughout the sample corpus and most real policy/compliance documents.
# Level = how many dot-separated numeric groups precede the heading text.
_NUMBERED_HEADING = re.compile(r"^(\d+(?:\.\d+)*)\.?\s+(\S.*)$")


class ExtractionError(Exception):
    """Raised for a corrupt or unreadable file — the pipeline turns this into
    a document.failed event (roadmap Phase 2 acceptance criterion 6)."""


def _heading_level_from_number(numbering: str) -> int:
    return numbering.count(".") + 1


_MAX_HEADING_LENGTH = 120


def _classify_line(line: str) -> tuple[bool, int | None]:
    stripped = line.strip()
    match = _NUMBERED_HEADING.match(stripped)
    if not match:
        return False, None

    title = match.group(2)
    # Distinguishes "1. Data Residency Requirements" (a heading) from
    # "15 business days prior to engagement." (ordinary prose that happens to
    # start with a number) — headings are short titles, not full sentences.
    looks_like_a_sentence = stripped.endswith((".", ",", ";")) or not title[:1].isupper()
    if looks_like_a_sentence or len(stripped) > _MAX_HEADING_LENGTH:
        return False, None

    return True, _heading_level_from_number(match.group(1))


def extract_pdf(path: Path) -> ExtractedDocument:
    blocks: list[ExtractedBlock] = []
    try:
        with pdfplumber.open(path) as pdf:
            for page_number, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                for raw_line in text.splitlines():
                    line = raw_line.strip()
                    if not line:
                        continue
                    is_heading, level = _classify_line(line)
                    blocks.append(
                        ExtractedBlock(
                            text=line,
                            page_number=page_number,
                            is_heading=is_heading,
                            heading_level=level,
                        )
                    )
    except Exception as e:
        raise ExtractionError(f"Failed to extract PDF: {e}") from e

    if not blocks:
        raise ExtractionError("PDF contains no extractable text (possibly scanned/image-only)")
    return ExtractedDocument(format="pdf", blocks=blocks)


def extract_docx(path: Path) -> ExtractedDocument:
    blocks: list[ExtractedBlock] = []
    try:
        doc = DocxDocument(str(path))
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = (paragraph.style.name if paragraph.style else "") or ""
            if style_name.startswith("Heading"):
                level_str = style_name.replace("Heading", "").strip()
                level: int | None = int(level_str) if level_str.isdigit() else 1
                blocks.append(ExtractedBlock(text=text, is_heading=True, heading_level=level))
            else:
                is_heading, level = _classify_line(text)
                blocks.append(ExtractedBlock(text=text, is_heading=is_heading, heading_level=level))
    except ExtractionError:
        raise
    except Exception as e:
        raise ExtractionError(f"Failed to extract DOCX: {e}") from e

    if not blocks:
        raise ExtractionError("DOCX contains no extractable text")
    return ExtractedDocument(format="docx", blocks=blocks)


def extract_txt(path: Path) -> ExtractedDocument:
    try:
        text = path.read_text(encoding="utf-8", errors="strict")
    except Exception as e:
        raise ExtractionError(f"Failed to extract TXT: {e}") from e

    blocks = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        is_heading, level = _classify_line(line)
        blocks.append(ExtractedBlock(text=line, is_heading=is_heading, heading_level=level))

    if not blocks:
        raise ExtractionError("TXT file is empty")
    return ExtractedDocument(format="txt", blocks=blocks)


_MD_HEADING = re.compile(r"^(#{1,6})\s+(\S.*)$")


def extract_md(path: Path) -> ExtractedDocument:
    try:
        text = path.read_text(encoding="utf-8", errors="strict")
    except Exception as e:
        raise ExtractionError(f"Failed to extract MD: {e}") from e

    blocks = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        md_match = _MD_HEADING.match(line)
        if md_match:
            level: int | None = len(md_match.group(1))
            blocks.append(
                ExtractedBlock(text=md_match.group(2), is_heading=True, heading_level=level)
            )
            continue
        is_heading, level = _classify_line(line)
        blocks.append(ExtractedBlock(text=line, is_heading=is_heading, heading_level=level))

    if not blocks:
        raise ExtractionError("MD file is empty")
    return ExtractedDocument(format="md", blocks=blocks)


_EXTRACTORS = {
    "pdf": extract_pdf,
    "docx": extract_docx,
    "txt": extract_txt,
    "md": extract_md,
}


def extract(path: Path, document_format: str) -> ExtractedDocument:
    extractor = _EXTRACTORS.get(document_format.lower())
    if extractor is None:
        raise ExtractionError(f"Unsupported format: {document_format}")
    return extractor(path)
